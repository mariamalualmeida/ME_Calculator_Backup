# file_io_utils.py

import pandas as pd
import io
import os
import re
from datetime import datetime

# As instalações de biblioteca serão no arquivo principal (main.py)
# Imports para o Code Interpreter (serão validados pelo main.py)
try:
    import pdfplumber
except ImportError:
    pass # Será instalado pelo main.py

try:
    from docx import Document
except ImportError:
    pass # Será instalado pelo main.py

try:
    import openpyxl
except ImportError:
    pass # Será instalado pelo main.py

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pass # Será instalado pelo main.py

try:
    import fitz # PyMuPDF
except ImportError:
    pass # Será instalado pelo main.py

try:
    from tabula import read_pdf
except ImportError:
    pass # Será instalado pelo main.py


def detect_file_type_by_filename(filename: str) -> str:
    """Detecta o tipo de documento (extrato, fatura, contracheque) pelo nome do arquivo."""
    filename_lower = filename.lower()
    if any(word in filename_lower for word in ['extrato', 'movimentacao', 'historico', 'comprovante']):
        return 'extrato_bancario'
    elif any(word in filename_lower for word in ['fatura', 'cartao', 'credito']):
        return 'fatura_cartao'
    elif any(word in filename_lower for word in ['contracheque', 'holerite', 'salario', 'renda']):
        return 'contracheque'
    return 'desconhecido'

def detect_bank_from_filename(filename: str) -> str:
    """Tenta detectar o banco pelo nome do arquivo."""
    filename_lower = filename.lower()
    if 'c6' in filename_lower:
        return 'C6 Bank'
    elif 'nubank' in filename_lower or 'nu_pagamentos' in filename_lower:
        return 'Nubank'
    elif 'caixa' in filename_lower:
        return 'Caixa Econômica Federal'
    elif 'inter' in filename_lower:
        return 'Banco Inter'
    elif 'picpay' in filename_lower:
        return 'PicPay'
    elif 'bradesco' in filename_lower:
        return 'Bradesco'
    elif 'itau' in filename_lower:
        return 'Itaú'
    elif 'santander' in filename_lower:
        return 'Santander'
    return 'Desconhecido'

def handle_uploaded_file(file_path: str, file_type: str) -> dict:
    """
    Lida com o upload e leitura de diferentes tipos de arquivos.
    Retorna texto e/ou DataFrames de tabelas.
    """
    extracted_text = ""
    extracted_tables = []

    if file_type == 'pdf':
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            try:
                                if table and len(table) > 1 and all(table[0]):
                                    df = pd.DataFrame(table[1:], columns=table[0])
                                else:
                                    df = pd.DataFrame(table)
                                extracted_tables.append(df)
                            except Exception as df_e:
                                print(f"Aviso: Não foi possível converter parte da tabela em DataFrame com pdfplumber: {df_e}. Extraindo como texto.")
                                extracted_text += "\n".join([str(item) for sublist in table for item in sublist if item is not None]) + "\n"
        except ImportError:
            print("pdfplumber não está instalado ou acessível. Não foi possível extrair texto/tabelas de PDF diretamente.")
        except Exception as e:
            print(f"Erro ao ler PDF com pdfplumber: {e}.")

    elif file_type == 'docx':
        try:
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            for table in doc.tables:
                data = []
                for i, row in enumerate(table.rows):
                    text = [cell.text for cell in row.cells]
                    data.append(text)
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0]) if len(data) > 1 and all(data[0]) else pd.DataFrame(data)
                    extracted_tables.append(df)
        except ImportError:
            print("python-docx não está instalado ou acessível. Não foi possível extrair texto/tabelas de DOCX.")
        except Exception as e:
            print(f"Erro ao ler DOCX: {e}")

    elif file_type == 'xlsx':
        try:
            import openpyxl # openpyxl é usado por pandas.read_excel
            xls = pd.ExcelFile(file_path)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                extracted_text += df.to_string(index=False) + "\n"
                extracted_tables.append(df)
        except ImportError:
            print("openpyxl não está instalado ou acessível. Não foi possível extrair dados de XLSX.")
        except Exception as e:
            print(f"Erro ao ler XLSX: {e}")

    elif file_type == 'csv':
        try:
            df = pd.read_csv(file_path)
            extracted_text = df.to_string(index=False)
            extracted_tables.append(df)
        except Exception as e:
            print(f"Erro ao ler CSV: {e}")

    elif file_type in ['jpg', 'png', 'jpeg']:
        pass # OCR is handled by perform_ocr in main processing flow

    return {"text": extracted_text, "tables": extracted_tables}

def perform_ocr(file_path_or_bytes: str | bytes) -> str:
    """
    Realiza OCR em um arquivo de imagem (JPG, PNG) ou PDF escaneado.
    :param file_path_or_bytes: Caminho do arquivo ou bytes (para arquivos em memória).
    :return: Texto extraído via OCR.
    """
    text = ""
    try:
        import pytesseract
        from PIL import Image
        import fitz # PyMuPDF

        if isinstance(file_path_or_bytes, str):
            if file_path_or_bytes.lower().endswith(('.jpg', '.jpeg', '.png')):
                img = Image.open(file_path_or_bytes)
                text = pytesseract.image_to_string(img, lang='por+eng')
            elif file_path_or_bytes.lower().endswith('.pdf'):
                pdf_document = fitz.open(file_path_or_bytes)
                for page_num in range(pdf_document.page_count):
                    page = pdf_document.load_page(page_num)
                    pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72)) # Aumentar DPI para melhor OCR
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text += pytesseract.image_to_string(img, lang='por+eng') + "\n"
                pdf_document.close()
        elif isinstance(file_path_or_bytes, bytes):
            img = Image.open(io.BytesIO(file_path_or_bytes))
            text = pytesseract.image_to_string(img, lang='por+eng')
    except ImportError:
        print("Erro: pytesseract, Pillow ou PyMuPDF não encontrados ou acessíveis. OCR indisponível.")
    except pytesseract.TesseractNotFoundError:
        print("Erro: Tesseract OCR não encontrado. Certifique-se de que está instalado no ambiente e configurado no PATH.")
    except Exception as e:
        print(f"Erro durante o OCR: {e}")
    return text